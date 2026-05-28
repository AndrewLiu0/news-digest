import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

def markdown_to_docx(markdown_text, output_path):
    """
    Converts the specific CCCW newsletter markdown format to a Word document.
    Handles bolding, underlining, and hyperlinks.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    lines = markdown_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        p = doc.add_paragraph()
        
        # Check for date headers (Underlined, e.g., __Friday, May 8, 2026__)
        date_header_match = re.match(r'^__(.*)__$', line)
        if date_header_match:
            run = p.add_run(date_header_match.group(1))
            run.bold = True
            run.underline = True
            continue

        # Parse links and bold text
        # Simple regex-based parser for [text](url) and **bold**
        parts = re.split(r'(\[.*?\]\(.*?\))|(\*\*.*?\*\*)', line)
        for part in parts:
            if not part:
                continue
            
            # Match Markdown Link: [text](url)
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            # Match Bold: **text**
            bold_match = re.match(r'\*\*(.*?)\*\*', part)
            
            if link_match:
                text = link_match.group(1)
                url = link_match.group(2)
                add_hyperlink(p, url, text)
            elif bold_match:
                text = bold_match.group(1)
                run = p.add_run(text)
                run.bold = True
            else:
                p.add_run(part)

    doc.save(output_path)

def add_hyperlink(paragraph, url, text):
    """
    Adds a hyperlink to a paragraph.
    """
    import docx
    # This is a bit of a hack as python-docx doesn't have a high-level API for hyperlinks
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Style the link
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0563C1') # Standard blue
    rPr.append(c)
    
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    text_obj = docx.oxml.shared.OxmlElement('w:t')
    text_obj.text = text
    new_run.append(text_obj)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

class CCCW_PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)
        self.add_page()
        self.set_font("Times", size=11)

    def write_markdown_line(self, line):
        line = line.strip()
        if not line:
            self.ln(5)
            return

        # Check for date headers
        date_header_match = re.match(r'^__(.*)__$', line)
        if date_header_match:
            self.set_font("Times", "BU", size=11)
            self.multi_cell(0, 6, date_header_match.group(1))
            self.set_font("Times", "", size=11)
            return

        # Parse links and bold text
        # This is simplified for PDF
        parts = re.split(r'(\[.*?\]\(.*?\))|(\*\*.*?\*\*)', line)
        
        current_x = self.get_x()
        for part in parts:
            if not part:
                continue
            
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            bold_match = re.match(r'\*\*(.*?)\*\*', part)
            
            if link_match:
                text = link_match.group(1)
                url = link_match.group(2)
                self.set_text_color(5, 99, 193) # Blue
                self.set_font("Times", "U", size=11)
                self.write(6, text, link=url)
                self.set_text_color(0, 0, 0)
                self.set_font("Times", "", size=11)
            elif bold_match:
                text = bold_match.group(1)
                self.set_font("Times", "B", size=11)
                self.write(6, text)
                self.set_font("Times", "", size=11)
            else:
                self.write(6, part)
        
        self.ln(6)

def clean_text(text):
    """
    Replaces common Unicode characters with ASCII equivalents for PDF generation.
    """
    replacements = {
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote
        '\u2013': '-',  # En dash
        '\u2014': '-',  # Em dash
        '\u2022': '*',  # Bullet
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Handle other potentially problematic characters by encoding and decoding
    return text.encode('ascii', 'ignore').decode('ascii')

def markdown_to_pdf(markdown_text, output_path):
    # Clean text before processing
    cleaned_markdown = clean_text(markdown_text)
    pdf = CCCW_PDF()
    lines = cleaned_markdown.split('\n')
    for line in lines:
        pdf.write_markdown_line(line)
    pdf.output(output_path)
